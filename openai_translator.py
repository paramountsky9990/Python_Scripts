import os
import openai
import docx
from docx.shared import Pt

# Replace with your actual API key
openai.api_key = "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"

from langdetect import detect

def is_english(text):
    try:
        detected_language = detect(text)
        return detected_language == 'en'
    except:
        return False

# Task 1: Japanese to English translation
def translate_japanese_to_english(text):
    prompt = f"Translate the following Japanese text to English: '{text}'"
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
                {"role": "user", "content": prompt}
            ]
    )    
    
    translated_text = response.choices[0].message.content.strip()
    if(is_english(translated_text)):
        print(f"output is {translated_text}")
        return translated_text
    

# Read Word document, translate content, and write to a new document
def translate_word_document(input_file, output_file):
    # Read the input Word document
    input_doc = docx.Document(input_file)
    # Create a new Word document for the translated content
    output_doc = docx.Document()

    for paragraph in input_doc.paragraphs:
        # Translate the paragraph
        translated_text = translate_japanese_to_english(paragraph.text)
        # Add the translated paragraph to the output document
        new_paragraph = output_doc.add_paragraph(translated_text)
        # Set the font size to match the original paragraph
        for run in paragraph.runs:
            font_size = run.font.size
            if font_size:
                new_run = new_paragraph.add_run(run.text)
                new_run.font.size = Pt(font_size)

    # Save the translated document
    output_doc.save(output_file)

# Example usage
input_file = "input.docx"
output_file = "output.docx"
translate_word_document(input_file, output_file)
